"""
generate_cv_docx.py — ATS-friendly Word CV generator.
Entry point: build_docx(json_path, docx_path)
Same JSON schema as generate_cv_last_version.py.
Optimised for Workday and other ATS parsers: no tables, no text boxes,
plain paragraphs, standard styles, extractable text.
"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para(doc, text="", bold=False, size=11, color=None,
          space_before=0, space_after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, color=color)
    return p


def _section_heading(doc, text, en=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, size=11, bold=True, color=(30, 30, 30))
    # Bottom border under heading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _entry_header(doc, title, period, company_location):
    """Role title | Period on one line, Company — Location on next."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)

    r1 = p.add_run(title)
    _set_font(r1, size=10, bold=True)

    if period:
        r2 = p.add_run(f"    {period}")
        _set_font(r2, size=10, bold=False, color=(80, 80, 80))

    if company_location:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r3 = p2.add_run(company_location)
        _set_font(r3, size=9.5, bold=False, color=(100, 100, 100))


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    _set_font(run, size=9.5)


def build_docx(json_path="cv_main.json", output_path="cv_output.docx"):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    en = data.get("_lang", "fr") == "en"
    identity = data["identity"]

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Header — Name ─────────────────────────────────────────────────────────
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    r = p_name.add_run(identity["full_name"])
    _set_font(r, size=20, bold=True)

    # Title
    display_title = data.get("_display_title") or identity.get("title", "")
    if display_title:
        _para(doc, display_title, size=12, color=(60, 60, 60), space_after=2)

    # Cert line
    display_cert = data.get("_display_cert", "")
    if display_cert:
        _para(doc, display_cert, size=10, color=(100, 100, 100), space_after=2)

    # Meta line (age · permis)
    meta_parts = [identity.get("age", ""), identity.get("permis", "")]
    meta_line = "   ·   ".join(p for p in meta_parts if p)
    if meta_line:
        _para(doc, meta_line, size=10, color=(120, 120, 120), space_after=2)

    # Summary
    summary = (data.get("profile") or {}).get("summary_fr", "")
    if summary:
        _para(doc, summary, size=10, color=(60, 60, 60), space_before=4, space_after=4)

    # Contact line — single paragraph, pipe-separated for clean ATS parsing
    import re
    linkedin_raw = identity.get("linkedin", "")
    linkedin = re.sub(r"https?://(www\.)?", "", linkedin_raw).rstrip("/")
    contact_parts = [
        identity.get("address", ""),
        identity.get("phone", ""),
        identity.get("email", ""),
        linkedin,
    ]
    contact_line = "   |   ".join(p for p in contact_parts if p)
    if contact_line:
        _para(doc, contact_line, size=9.5, color=(80, 80, 80), space_after=6)

    # ── Experience ────────────────────────────────────────────────────────────
    _section_heading(doc, "Experience" if en else "Expériences")
    for exp in data.get("experience", []):
        company_loc = f"{exp['company']} — {exp['location']}"
        _entry_header(doc, exp["title"], exp.get("period", ""), company_loc)
        for bullet in exp.get("bullets", []):
            _bullet(doc, bullet)

    # ── Education ─────────────────────────────────────────────────────────────
    _section_heading(doc, "Education" if en else "Études — Diplômes")
    for edu in data.get("education", []):
        school_loc = f"{edu['school']} — {edu['location']}"
        _entry_header(doc, edu["degree"], edu.get("year", ""), school_loc)
        if edu.get("details"):
            _bullet(doc, edu["details"])

    # ── Projects ──────────────────────────────────────────────────────────────
    _section_heading(doc, "Personal Projects" if en else "Projets personnels")
    for proj in data.get("projects", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        r_name = p.add_run(proj["name"])
        _set_font(r_name, size=10, bold=True)
        tech = proj.get("tech", [])
        if tech:
            r_tech = p.add_run(f"   ({', '.join(tech)})")
            _set_font(r_tech, size=9.5, color=(100, 100, 100))
        if proj.get("description"):
            _bullet(doc, proj["description"])

    # ── Skills ────────────────────────────────────────────────────────────────
    _section_heading(doc, "Skills" if en else "Compétences")
    skills = data.get("skills", [])
    # Group by label
    groups: dict = {}
    for sk in skills:
        groups.setdefault(sk["label"], []).append(sk["value"])
    for label, values in groups.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r_lbl = p.add_run(f"{label} : ")
        _set_font(r_lbl, size=9.5, bold=True)
        r_val = p.add_run(", ".join(values))
        _set_font(r_val, size=9.5)

    # ── Certifications ────────────────────────────────────────────────────────
    certs = data.get("certifications", [])
    if certs:
        _section_heading(doc, "Certifications")
        for cert in certs:
            _bullet(doc, f"{cert['name']} ({cert['status']}) — {cert['issuer']}")

    # ── Languages ─────────────────────────────────────────────────────────────
    langs = data.get("languages", [])
    if langs:
        _section_heading(doc, "Languages" if en else "Langues")
        for lang in langs:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r_lbl = p.add_run(f"{lang['label']} : ")
            _set_font(r_lbl, size=9.5, bold=True)
            r_val = p.add_run(lang["value"])
            _set_font(r_val, size=9.5)

    doc.save(output_path)
    print(f"Generated: {output_path}")


if __name__ == "__main__":
    import sys
    json_in  = sys.argv[1] if len(sys.argv) > 1 else "cv_main.json"
    docx_out = sys.argv[2] if len(sys.argv) > 2 else "cv_output.docx"
    build_docx(json_in, docx_out)
